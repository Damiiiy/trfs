# from docx import Document
# import re
# from .models import *

# def import_questions_from_docx(file, exam_instance):
#     doc = Document(file)
#     questions = []
#     current_question = {}
#     question_pattern = re.compile(r'^\d+\.\s+(.*)')
#     option_pattern = re.compile(r'^([A-D])\.\s+(.*?)(✅)?$')

#     for para in doc.paragraphs:
#         text = para.text.strip()
#         if not text:
#             continue

#         q_match = question_pattern.match(text)
#         o_match = option_pattern.match(text)

#         if q_match:
#             if current_question:
#                 questions.append(current_question)
#             current_question = {
#                 'question_text': q_match.group(1),
#                 'options': {},
#                 'correct': None
#             }

#         elif o_match and current_question:
#             label = o_match.group(1)
#             option_text = o_match.group(2)
#             is_correct = o_match.group(3)
#             current_question['options'][label] = option_text
#             if is_correct:
#                 current_question['correct'] = label

#     if current_question:
#         questions.append(current_question)

#     # Save to DB
#     for q in questions:
#         correct_map = {
#             'A': 'A',
#             'B': 'B',
#             'C': 'C',
#             'D': 'D'
#         }
#         Question.objects.create(
#             exams=exam_instance,
#             question=q['question_text'],
#             option1=q['options'].get('A', ''),
#             option2=q['options'].get('B', ''),
#             option3=q['options'].get('C', ''),
#             option4=q['options'].get('D', ''),
#             answer=correct_map.get(q['correct'], 'Option1')  # Default fallback
#         )


from docx import Document
import re
from .models import Question

def import_questions_from_docx(file, exam_instance):
    doc = Document(file)
    questions = []
    current_question = {}
    
    # Regex patterns
    question_pattern = re.compile(r'^\d+\.\s+(.*)')
    option_pattern = re.compile(r'^([A-D])\.\s+(.*?)(✅)?$')

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Match question line
        q_match = question_pattern.match(text)
        if q_match:
            if current_question:
                questions.append(current_question)
            current_question = {
                'question_text': q_match.group(1),
                'options': {},
                'correct': None
            }
            continue

        # Match option line
        o_match = option_pattern.match(text)
        if o_match and current_question:
            label = o_match.group(1)  # 'A', 'B', 'C', or 'D'
            option_text = o_match.group(2)
            is_correct = o_match.group(3)
            current_question['options'][label] = option_text
            if is_correct:
                current_question['correct'] = label

    # Append last question
    if current_question:
        questions.append(current_question)

    # Save to database
    for q in questions:
        Question.objects.create(
            exams=exam_instance,
            question=q['question_text'],
            option1=q['options'].get('A', ''),
            option2=q['options'].get('B', ''),
            option3=q['options'].get('C', ''),
            option4=q['options'].get('D', ''),
            answer=q['correct'] or 'A'  # Default to 'A' if no correct answer marked
        )
